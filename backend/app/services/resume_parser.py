"""
Resume text extraction.

Supported now:  PDF (.pdf), Word (.docx / .doc)
Scaffolded:     Image OCR (.jpg/.jpeg/.png) — hook is here, implementation deferred
Stubbed:        Job-board import (Naukri/LinkedIn) — needs paid API access
"""
import io
import os


def extract_text_from_pdf(file_bytes: bytes) -> str:
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    return "\n".join(
        page.extract_text() or "" for page in reader.pages
    ).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    # Paragraphs (body text, headings)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Table cells — names/contact info are often in styled table headers
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in parts:
                    parts.append(text)
    return "\n".join(parts)


def extract_text_from_image(file_bytes: bytes) -> str:
    # OCR hook — deferred to a later sub-step.
    # Requires pytesseract + Tesseract binary, or GCP Vision API.
    raise NotImplementedError(
        "Image OCR is not yet implemented. "
        "Upload a PDF or Word document instead."
    )


def import_from_jobboard(source: str, profile_url: str) -> dict:
    # Naukri / LinkedIn import stub.
    # Needs a paid API agreement or partnership with the job board.
    raise NotImplementedError(
        f"Import from '{source}' is not yet supported. "
        "This requires a paid API agreement with the job board."
    )


def normalize_phone(phone: str) -> str | None:
    """
    Strip formatting and country codes from a phone number, returning 10 digits
    for Indian mobiles or at least 8 digits for others.  Returns None if the
    input is blank or produces fewer than 8 digits.
    """
    import re
    if not phone:
        return None
    digits = re.sub(r'[^\d]', '', phone)
    # +91 / 0 prefix → strip to 10-digit Indian mobile
    if len(digits) == 12 and digits.startswith('91'):
        digits = digits[2:]
    elif len(digits) == 11 and digits.startswith('0'):
        digits = digits[1:]
    return digits if len(digits) >= 8 else None


def extract_contact_info(text: str) -> dict:
    """
    Heuristic extraction of name, email, and phone from raw resume text.
    Used to pre-fill the submit-application form.
    """
    import re

    # Email — reliable
    email_match = re.search(r'[\w.+\-]+@[\w\-]+\.[a-zA-Z]{2,}', text)
    email = email_match.group().lower() if email_match else None

    # Phone — tries Indian mobile (6–9 start, 10 digits) first, then generic
    phone_raw = None
    # Indian mobile: optional +91 / 0 prefix, then 10 digits starting 6-9
    m = re.search(r'(?:(?:\+91|0)[\s\-]?)?([6-9]\d{2}[\s\-]?\d{3}[\s\-]?\d{4})', text)
    if m:
        phone_raw = m.group(0).strip()
    else:
        # Generic international: +CC (nnn) nnn-nnnn style
        m2 = re.search(r'(?:\+\d{1,3}[\s\-]?)?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}', text)
        if m2:
            phone_raw = m2.group(0).strip()

    phone = normalize_phone(phone_raw) if phone_raw else None
    # Return the original formatted string for display; backend normalizes for dedup
    phone_display = phone_raw if phone else None

    # Name — first short alphabetic line that isn't a section header or label
    _SKIP_WORDS = {
        "curriculum", "vitae", "resume", "cv", "profile", "objective",
        "summary", "education", "experience", "skills", "contact",
        "address", "declaration", "references", "achievements",
        "personal", "details", "information", "hobbies", "interests",
    }
    name = None
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped or len(stripped) > 60:
            continue
        if re.search(r'[\d@#/\\|<>{}:;]', stripped):
            continue
        words = stripped.split()
        if not (2 <= len(words) <= 5):
            continue
        if not all(re.match(r"[A-Za-z'\-\.]+$", w) for w in words):
            continue
        # Skip lines that are just section headings
        lower_words = {w.lower() for w in words}
        if lower_words & _SKIP_WORDS:
            continue
        name = stripped.title()
        break

    # Current company — compact line containing a known corporate suffix
    _co_suffix = re.compile(
        r'\b(pvt\.?|ltd\.?|limited|inc\.?|corp\.?|llc|technologies|solutions|'
        r'systems|services|software|tech|group|holdings|enterprises|consulting)\b',
        re.IGNORECASE,
    )
    # Current designation — compact line containing a known job title keyword
    _title_kw = re.compile(
        r'\b(engineer|developer|analyst|manager|architect|consultant|director|lead|'
        r'head|officer|executive|specialist|coordinator|associate|programmer|designer|'
        r'scientist|administrator|intern|trainee)\b',
        re.IGNORECASE,
    )
    _section_words = {
        'experience', 'education', 'skills', 'summary', 'objective', 'profile',
        'employment', 'career', 'academic', 'qualification', 'certification',
    }
    current_company = None
    current_designation = None
    for line in (text or "").splitlines()[:60]:
        stripped = line.strip()
        if not stripped or len(stripped) > 100 or len(stripped) < 4:
            continue
        if re.search(r'[@/\\|<>{}:;]|http|www\.', stripped):
            continue
        words = stripped.split()
        if not (1 <= len(words) <= 10):
            continue
        if name and stripped.lower() == name.lower():
            continue
        lower = stripped.lower()
        # Skip bare section headers
        if any(sw in lower for sw in _section_words) and len(words) <= 2:
            continue
        if current_designation is None and _title_kw.search(stripped):
            current_designation = stripped
        if current_company is None and _co_suffix.search(stripped) and not _title_kw.search(stripped):
            current_company = stripped

    return {
        "full_name": name,
        "email": email,
        "phone": phone_display,
        "current_company": current_company,
        "current_designation": current_designation,
    }


def extract_text(file_bytes: bytes, filename: str) -> tuple:
    """
    Extract resume text from an uploaded file.
    Returns (text: str, warning: str).
    `warning` is empty on clean success; non-empty if extraction was partial or
    the file type fell through to a stub.
    """
    suffix = os.path.splitext(filename)[1].lower()

    if suffix == ".pdf":
        try:
            text = extract_text_from_pdf(file_bytes)
            if not text:
                return "", "PDF parsed but no text found — it may be a scanned/image-only PDF."
            return text, ""
        except Exception as exc:
            return "", f"PDF parsing error: {exc}"

    if suffix in (".docx", ".doc"):
        try:
            text = extract_text_from_docx(file_bytes)
            if not text:
                return "", "Word document parsed but contained no readable text."
            return text, ""
        except Exception as exc:
            return "", f"Word document parsing error: {exc}"

    if suffix in (".jpg", ".jpeg", ".png"):
        raise NotImplementedError(
            "Image OCR is not yet implemented. "
            "Upload a PDF or Word document instead."
        )

    return "", f"Unsupported file type '{suffix}'."
