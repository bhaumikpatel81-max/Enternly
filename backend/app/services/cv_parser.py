"""
CV / Resume parser.

Tier-1 (synchronous, instant):
  - sha256 dedup hash
  - Text extraction: PDF via PyMuPDF, DOCX via python-docx (mammoth fallback),
    DOC via LibreOffice headless conversion then DOCX parse.
  - Candidate name from filename convention.
  - Skills: word-boundary match against skills_dictionary.json.
  - tsvector generation SQL fragment (done inside the INSERT, not here).

Tier-2 (async, rate-limited) is handled separately by cv_enricher.py.
"""
import hashlib
import io
import json
import os
import re
import shutil
import subprocess
import tempfile
import threading
from pathlib import Path
from typing import Optional

# ── Concurrency caps for CPU-heavy extraction ────────────────────────────────
# Scans can now overlap freely (the automatic per-recruiter poller runs every
# few minutes independent of any manual "Scan My Email"/"Scan Ingest Folder"
# trigger, and nothing stops several recruiters from triggering their own
# scans at the same time). Each attachment's text extraction is genuine
# CPU-bound work — PyMuPDF/pypdf parsing, and .doc files spawn a full
# LibreOffice process — so with no cap, enough of these overlapping at once
# on a small box pegs every core just from extraction, on top of whatever
# else the server is doing (including answering health checks). These
# threading.Semaphores (not asyncio.Semaphore — this code runs inside
# worker threads via asyncio.to_thread/BackgroundTasks, not directly on the
# event loop) throttle how many extractions run at once system-wide;
# anything beyond the cap just waits its turn instead of piling on.
_EXTRACT_CONCURRENCY      = int(os.environ.get("CV_EXTRACT_CONCURRENCY", "3"))
_LIBREOFFICE_CONCURRENCY  = int(os.environ.get("CV_LIBREOFFICE_CONCURRENCY", "1"))
_extract_semaphore     = threading.Semaphore(_EXTRACT_CONCURRENCY)
_libreoffice_semaphore = threading.Semaphore(_LIBREOFFICE_CONCURRENCY)

# ── Skills dictionary (loaded once) ──────────────────────────────────────────

_SKILLS_PATH = Path(__file__).resolve().parent.parent / "data" / "skills_dictionary.json"
_SKILLS_RE: Optional[re.Pattern] = None

def _load_skills_re() -> re.Pattern:
    global _SKILLS_RE
    if _SKILLS_RE is None:
        with open(_SKILLS_PATH, encoding="utf-8") as f:
            data = json.load(f)
        skills = sorted(data["skills"], key=len, reverse=True)  # longer first
        # Word boundary; allow hyphens/underscores inside skill names
        parts = [re.escape(s).replace(r"\_", r"[_\-]").replace(r"\+", r"\+") for s in skills]
        pattern = r'(?<![a-zA-Z0-9_\-])(' + '|'.join(parts) + r')(?![a-zA-Z0-9_\-])'
        _SKILLS_RE = re.compile(pattern, re.IGNORECASE)
    return _SKILLS_RE


# ── Hash ──────────────────────────────────────────────────────────────────────

def sha256_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# ── Name from filename ────────────────────────────────────────────────────────

def parse_candidate_name(filename: str) -> Optional[str]:
    """
    Extract a candidate name from a filename.

    Rules (applied in order):
      1. Strip extension.
      2. Split on '_Resume' or '_resume' (case-insensitive); take the part BEFORE.
         e.g. 'John_Smith_Resume.pdf' → 'John Smith'
      3. Replace remaining underscores/hyphens with spaces, strip, title-case.
      4. If result is empty or looks like a non-name (all digits, single char), return None.
    """
    stem = Path(filename).stem  # no extension
    # Try splitting on _Resume / -Resume / _CV / -CV / _cv
    m = re.split(r'[_\-](?:resume|cv|curriculum|vitae)\b', stem, maxsplit=1, flags=re.IGNORECASE)
    name_part = m[0] if len(m) > 1 else stem

    # Replace underscores, hyphens, multiple spaces with single space
    name = re.sub(r'[_\-]+', ' ', name_part)
    name = re.sub(r'\s+', ' ', name).strip()

    # Reject if empty, too short, or all digits/non-alpha
    if not name or len(name) < 3 or not re.search(r'[A-Za-z]{2,}', name):
        return None

    return name.title()


# ── Text extraction ───────────────────────────────────────────────────────────

def _extract_pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=data, filetype="pdf")
        try:
            pages = [doc[i].get_text() for i in range(doc.page_count)]
            return "\n".join(pages).strip()
        finally:
            doc.close()
    except Exception:
        # Fallback to pypdf if fitz unavailable
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages).strip()


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t and t not in parts:
                        parts.append(t)
        text = "\n".join(parts).strip()
        if text:
            return text
    except Exception:
        pass
    # mammoth fallback — extracts as plain text from docx
    try:
        import mammoth
        result = mammoth.extract_raw_text(io.BytesIO(data))
        return (result.value or "").strip()
    except Exception:
        pass
    return ""


def _extract_doc(data: bytes) -> str:
    """Convert .doc → .docx via LibreOffice headless, then parse as docx."""
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return ""
    # Spawning a full LibreOffice process is by far the heaviest single
    # operation in this module — capped separately (and more tightly) than
    # the general extraction semaphore below so a batch of legacy .doc
    # files can't launch a pile of office-suite processes at once.
    with _libreoffice_semaphore:
        with tempfile.TemporaryDirectory() as tmpdir:
            src = os.path.join(tmpdir, "input.doc")
            with open(src, "wb") as f:
                f.write(data)
            try:
                subprocess.run(
                    [soffice, "--headless", "--convert-to", "docx",
                     "--outdir", tmpdir, src],
                    timeout=30,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    check=True,
                )
            except Exception:
                return ""
            out_path = os.path.join(tmpdir, "input.docx")
            if not os.path.exists(out_path):
                return ""
            with open(out_path, "rb") as f:
                return _extract_docx(f.read())


def extract_text(data: bytes, ext: str) -> str:
    """
    Extract plaintext from CV bytes given the file extension (.pdf/.docx/.doc).
    Returns empty string on failure — callers decide how to handle.

    Throttled by _extract_semaphore — real CPU-bound parsing work, capped so
    overlapping scans (the automatic mailbox poller, manual scans, and
    uploads can all run at the same time) don't collectively saturate every
    core at once.
    """
    ext = ext.lower().lstrip(".")
    if ext not in ("pdf", "docx", "doc", "txt"):
        return ""
    with _extract_semaphore:
        if ext == "pdf":
            return _extract_pdf(data)
        if ext == "docx":
            return _extract_docx(data)
        if ext == "doc":
            return _extract_doc(data)
        return data.decode("utf-8", errors="ignore").strip()


# ── Parse quality assessment (Improvement 4) ─────────────────────────────────

def assess_parse_quality(text: str, file_size_bytes: int = 0) -> dict:
    """
    Detect whether extracted text represents a parseable (text-based) resume
    or is likely image-based / scanned (low character yield per KB).

    Returns a dict ready to merge into score_breakdown JSONB.
    Does NOT reject candidates — flags only.
    """
    char_count = len((text or "").strip())

    if file_size_bytes > 0:
        file_size_kb = file_size_bytes / 1024
    else:
        # Estimate: plain-text resumes typically run ~2 KB per 1000 chars
        file_size_kb = max(char_count / 1000, 1)

    parseability_ratio = char_count / max(file_size_kb, 1)

    if parseability_ratio < 10 or char_count < 100:
        parse_quality = "low"
    elif char_count < 500:
        parse_quality = "medium"
    else:
        parse_quality = "good"

    return {
        "parse_quality":             parse_quality,
        "char_count":                char_count,
        "manual_review_recommended": parse_quality == "low",
    }


# ── CV-vs-not-a-CV content classification ────────────────────────────────────
# Used by any bulk ingestion path that can't rely on a human having
# curated the file list — email attachments (cv_email_scan.py) and the
# "Scan Ingest Folder" bulk_folder path (cv_api.py) both hand every
# PDF/DOCX/DOC they find here before storing it, since neither one can
# assume every file dropped in that folder/inbox is actually a resume
# (ID scans, payslips, bank statements, offer letters, plain photos all
# show up in real inboxes/folders alongside genuine CVs).

_RESUME_POSITIVE_KEYWORDS = [
    "experience", "education", "skills", "career summary",
    "professional summary", "employment history", "work history",
    "certification", "certifications", "projects", "qualification",
    "curriculum vitae", "resume", "profile summary", "achievements",
    "references", "career history", "personal details",
    "technical skills", "key skills", "core competencies",
    "linkedin", "github", "internship",
    "professional experience", "work experience", "career objective",
]

_RESUME_NEGATIVE_KEYWORDS = [
    "invoice", "gstin", "purchase order", "bill to", "amount due",
    "tax invoice", "terms of employment", "offer of employment",
    "offer letter", "non-disclosure agreement", "salary slip", "payslip",
    "bank statement", "aadhaar", "passport no", "boarding pass",
    "invoice number", "quotation", "credit note", "appointment letter",
    "increment letter", "relieving letter", "experience letter",
    "medical fitness", "fitness declaration", "medical certificate",
    "fit to work", "cover letter",
    # Job descriptions read deceptively like resumes — same vocabulary
    # (experience, skills, qualifications, responsibilities) describing a
    # role's requirements instead of one person's actual background.
    "job description", "job purpose", "key result area", "kra",
    "roles and responsibilities", "reporting to", "reports to",
    "we are looking for", "the ideal candidate",
]

# A resume needs at least this many DISTINCT positive keywords to be kept —
# a single generic word (e.g. "declaration", "objective") can show up in
# all sorts of non-resume documents; requiring 2+ cuts false positives like
# a medical fitness certificate that happens to mention "declaration" once.
_MIN_POSITIVE_HITS = 2


_FILENAME_RESUME_HINT_RE = re.compile(r'(resume|cv|curriculum|vitae)', re.IGNORECASE)

# A genuine resume almost always lists the candidate's own contact details
# somewhere; a JD, SOP, offer letter, or similar document usually doesn't
# carry a specific person's personal email/phone even when its wording
# happens to overlap with resume vocabulary (experience, skills,
# qualifications). Requiring one of these as well as keyword signals
# catches documents that read resume-ish but belong to no one in particular.
_EMAIL_RE = re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}')
_PHONE_RE = re.compile(r'(?:\+?\d[\d\-\s().]{7,16}\d)')

# Terms that essentially never appear in a genuine resume's own filename —
# unlike body-text keywords, a filename explicitly naming the document type
# ("Offer Letter ....pdf", "PAN Card.pdf") is close to unambiguous, so this
# rejects on filename alone before even looking at the extracted text.
_FILENAME_NEGATIVE_HINT_RE = re.compile(
    r'(offer[\s_-]*letter|appointment[\s_-]*letter|relieving[\s_-]*letter|increment[\s_-]*letter|'
    r'experience[\s_-]*letter|cover[\s_-]*letter|salary[\s_-]*slip|pay[\s_-]*slip|bank[\s_-]*statement|'
    r'pan[\s_-]*card|aadhaar|passport|invoice|quotation|boarding[\s_-]*pass|'
    r'medical[\s_-]*(fitness|certificate)|fitness[\s_-]*declaration|'
    r'\bjd\b|job[\s_-]*description|job[\s_-]*desc\b)',
    re.IGNORECASE,
)


def classify_resume_text(text: str, filename: str = "") -> tuple[bool, str]:
    """
    Decide whether an attachment's extracted text looks like an actual
    resume/CV, vs. some other document (offer letter, invoice, ID scan,
    payslip, photo, bank statement, a generic "Download.pdf", etc.) that
    happened to also be a PDF/DOC/DOCX. Used to filter multi-attachment
    emails (e.g. a Naukri response that includes a resume plus an
    unrelated form) and folder dumps (a "CV inbox" that in practice also
    collects ID scans, payslips, offer letters, and random reports) down
    to just the genuine CVs.

    Precision over recall: a resume must show a positive signal (keyword
    match, or an explicit resume/CV/curriculum-vitae filename) to be kept.
    Anything ambiguous — no resume content, no negative match either, just
    generic text — is rejected rather than defaulted to "keep", since a
    wrongly-rejected file just needs a manual re-upload but a wrongly-kept
    one pollutes the repository silently.

    Returns (is_resume, reason).
    """
    t = (text or "").lower().strip()
    fname_hint = bool(_FILENAME_RESUME_HINT_RE.search(filename or ""))

    # Filename says "Offer Letter", "PAN Card", "Salary Slip", etc. — trust
    # that over content keywords, which can overlap with resume wording
    # (an offer letter mentioning "your experience and qualifications"
    # would otherwise slip past a purely content-based check).
    if _FILENAME_NEGATIVE_HINT_RE.search(filename or "") and not fname_hint:
        return False, "filename_names_a_non_resume_document"

    if not t:
        # No extractable text at all (a scanned/image-only file, or a
        # non-document image saved with a .pdf/.docx extension).
        if fname_hint:
            return True, "no_extractable_text_but_filename_suggests_resume"
        return False, "no_extractable_text_and_no_resume_filename_hint"

    neg_hits = sum(1 for k in _RESUME_NEGATIVE_KEYWORDS if k in t)
    pos_hits = sum(1 for k in _RESUME_POSITIVE_KEYWORDS if k in t)

    # Negative signals win over positive ones (unless the filename itself
    # explicitly says resume/CV) — a job description or offer letter easily
    # racks up 2+ "positive" hits too (experience/skills/qualifications are
    # exactly the words a JD uses to describe a ROLE's requirements, not a
    # person's actual background), so letting positive count win first let
    # real JDs through undetected. A generic-named document that clearly
    # reads like a JD/offer-letter/etc. should lose even if it also
    # mentions resume-ish words.
    if neg_hits >= 1 and not fname_hint:
        return False, "matches_non_resume_document_pattern"

    has_contact = bool(_EMAIL_RE.search(t) or _PHONE_RE.search(t))

    if pos_hits >= _MIN_POSITIVE_HITS:
        if has_contact or fname_hint:
            return True, "resume_keywords_found"
        # Reads resume-ish (experience/skills/education sections) but no
        # candidate email or phone anywhere in it — exactly the profile of
        # a JD/SOP/policy document that happens to share resume vocabulary
        # without belonging to any one person. Filename hint above already
        # covers the legitimate exception (a real resume genuinely missing
        # contact info in its extracted text but explicitly named Resume/CV).
        return False, "resume_keywords_but_no_candidate_contact_info"
    if pos_hits >= 1 and fname_hint:
        # One keyword plus an explicit resume/CV filename together are
        # enough even though one keyword alone isn't.
        return True, "resume_keyword_and_filename_hint"
    if fname_hint:
        return True, "no_resume_keywords_but_filename_suggests_resume"
    return False, "no_resume_signal_found"


# ── Tier-1 skill extraction ───────────────────────────────────────────────────

def extract_tier1_skills(raw_text: str) -> list[str]:
    """
    Fast keyword match against skills_dictionary.json.
    Returns sorted, deduplicated lowercase skill list.
    """
    if not raw_text:
        return []
    pattern = _load_skills_re()
    matches = pattern.findall(raw_text)
    return sorted({m.lower().replace("-", "_") for m in matches})
